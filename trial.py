import json
from fastapi import FastAPI, Form, HTTPException, Request, Depends, UploadFile, File,APIRouter
from fastapi.templating import Jinja2Templates
from sqlalchemy import create_engine, Column, Integer, String, Text, ForeignKey,JSON
from sqlalchemy.ext.declarative import declarative_base
from sqlalchemy.orm import sessionmaker, relationship, Session
from passlib.context import CryptContext
from pydantic_ai import Agent
from pydantic_ai.models.groq import GroqModel
from fastapi.responses import HTMLResponse, JSONResponse, RedirectResponse,StreamingResponse
from pydantic import BaseModel
from starlette.middleware.sessions import SessionMiddleware
import os
import chromadb
from chromadb.config import Settings
from fastapi.staticfiles import StaticFiles
import uuid
from datetime import datetime
from io import BytesIO
from docx import Document as DocxDocument
from docx.shared import Inches
from reportlab.lib import colors
from reportlab.lib.pagesizes import letter
from reportlab.pdfgen import canvas
from reportlab.platypus import SimpleDocTemplate, Table, TableStyle, Paragraph
from reportlab.lib.styles import getSampleStyleSheet
from PyPDF2 import PdfReader

DATABASE_URL = "mysql+pymysql://root:Dhruvil%402003@localhost/project_management"
engine = create_engine(DATABASE_URL, pool_recycle=3600, pool_size=10)
SessionLocal = sessionmaker(autocommit=False, autoflush=False, bind=engine)

Base = declarative_base()

class User(Base):
    __tablename__ = 'users'
    id = Column(Integer, primary_key=True, index=True)
    first_name = Column(String(50))
    last_name = Column(String(50))
    dob = Column(String(10))
    country = Column(String(50))
    country_code = Column(String(10))
    phone_number = Column(String(20), unique=True, index=True)
    gender = Column(String(10))
    email = Column(String(100), unique=True, index=True)
    password = Column(String(255))
    profile_image = Column(String(255)) 

# Password hashing setup
pwd_context = CryptContext(schemes=["bcrypt"], deprecated="auto")

class Epic(Base):
    __tablename__ = 'epic'
    id = Column(Integer, primary_key=True, index=True)
    user_id = Column(Integer, ForeignKey('users.id'), nullable=True) 
    title = Column(String(255), index=True)
    description = Column(Text)
    estimation = Column(Integer)
    stories = relationship('Story', back_populates='epic', cascade='all, delete-orphan')
    user = relationship('User')

class Story(Base):
    __tablename__ = 'stories'
    id = Column(Integer, primary_key=True, index=True)
    user_id = Column(Integer, ForeignKey('users.id'), nullable=True)  # Changed to nullable
    epic_id = Column(Integer, ForeignKey('epic.id'))
    title = Column(String(255), index=True)
    description = Column(Text)
    estimation = Column(Integer)
    tasks = relationship('Task', back_populates='story', cascade='all, delete-orphan')
    epic = relationship('Epic', back_populates='stories')
    user = relationship('User')

class Task(Base):
    __tablename__ = 'tasks'
    id = Column(Integer, primary_key=True, index=True)
    user_id = Column(Integer, ForeignKey('users.id'), nullable=True)  # Changed to nullable
    story_id = Column(Integer, ForeignKey('stories.id'))
    title = Column(String(255), index=True)
    description = Column(Text)
    estimation = Column(Integer)
    story = relationship('Story', back_populates='tasks')
    user = relationship('User')

class Chat(Base):
    __tablename__ = 'chats'
    id = Column(Integer, primary_key=True, index=True)
    user_id = Column(Integer, ForeignKey('users.id'))
    chat_name = Column(String(255))
    created_at = Column(String(20), default=datetime.now().strftime("%Y-%m-%d %H:%M:%S"))
    conversations = relationship('Conversation', back_populates='chat', cascade='all, delete-orphan')

class Conversation(Base):
    __tablename__ = 'conversations'
    id = Column(Integer, primary_key=True, index=True)
    chat_id = Column(Integer, ForeignKey('chats.id'))
    user_message = Column(Text)
    bot_response = Column(Text)
    structured_data = Column(JSON)
    created_at = Column(String(20), default=datetime.now().strftime("%Y-%m-%d %H:%M:%S"))
    chat = relationship('Chat', back_populates='conversations')

class TaskBase(BaseModel):
    title: str
    description: str
    estimation: int

class StoryBase(BaseModel):
    title: str
    description: str
    estimation: int
    tasks: list[TaskBase]

class EpicBase(BaseModel):
    title: str
    description: str
    estimation: int
    stories: list[StoryBase]

class AIResponse(BaseModel):
    epics: list[EpicBase]

system_prompt = '''Provide answer in epic, stories, and tasks only. Make sure that every epic should have multiple stories and every story should have multiple tasks. Every epic, story, and task must have a title, short description (1 to 2 lines) and estimation (how many hours will be required to achieve that thing and it should be practical). Provide at least 3 epics.'''

app = FastAPI()

SECRET_KEY = os.urandom(32).hex()
app.add_middleware(SessionMiddleware, secret_key=SECRET_KEY, session_cookie="myapp_session")

chroma_client = chromadb.Client()
collection = chroma_client.get_or_create_collection(name="epics_stories_tasks")

model = GroqModel('llama-3.3-70b-versatile', api_key='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')
agent = Agent(model, system_prompt=system_prompt, result_type=AIResponse)

async def generate_formatted_response(answer: str) -> str:
    prompt = '''Format:
                epic 1: [Epic Title]
                    description: [Short Description of Epic]
                    estimation: [Number of hours for Epic]
                        story 1: [Story Title]
                            description: [Short Description of Story]
                            estimation: [Number of hours for Story]
                                task 1: [Task Title]
                                    description: [Short Description of Task]
                                    estimation: [Number of hours for Task]
                                task 2: [Task Title]
                                    description: [Short Description of Task]
                                    estimation: [Number of hours for Task]
                        story 2: [Story Title]
                            description: [Short Description of Story]
                            estimation: [Number of hours for Story]
                                task 1: [Task Title]
                                    description: [Short Description of Task]
                                    estimation: [Number of hours for Task]
                                task 2: [Task Title]
                                    description: [Short Description of Task]
                                    estimation: [Number of hours for Task]

'''
    agent_a = Agent(model, system_prompt=prompt)
    formatted_output = await agent_a.run(f"Format this output: {answer}")
    return formatted_output.data

templates = Jinja2Templates(directory=".")

def get_db():
    db = SessionLocal()
    return db

UPLOAD_DIR = "uploads"
os.makedirs(UPLOAD_DIR, exist_ok=True)
app.mount("/uploads", StaticFiles(directory="uploads"), name="uploads")

@app.get("/file-answer.html", response_class=HTMLResponse)
async def file_answer_page(request: Request):
    return templates.TemplateResponse("file-answer.html", {"request": request})

@app.post("/submit-file-followup")
async def submit_file_followup(
    request: Request,
    db: Session = Depends(get_db)
):
    try:
        data = await request.json()
        question = data.get("question")
        previous_output = data.get("previous_output")
        
        if not question or not previous_output:
            raise HTTPException(status_code=400, detail="Question and previous output are required")

        prompt = f"Previous response:\n{previous_output}\n\nNew request: {question}\n\n"
        prompt += "Please modify or add to the previous response while maintaining the same format."

        result = await agent.run(prompt)
        answer = result.data
        formatted_output = await generate_formatted_response(str(answer))

        return JSONResponse({
            "formatted_output": formatted_output,
            "answer": answer.dict()
        })

    except Exception as e:
        print(f"Error in submit_file_followup: {str(e)}")
        raise HTTPException(status_code=500, detail=str(e))

@app.post("/upload-pdf-query")
async def upload_pdf_query(request: Request, file: UploadFile = File(...), db: Session = Depends(get_db)):
    user_email = request.session.get("user", {}).get("email")
    if not user_email:
        raise HTTPException(status_code=401, detail="User not logged in")
    
    user = db.query(User).filter(User.email == user_email).first()
    if not user:
        raise HTTPException(status_code=404, detail="User not found")

    filename = file.filename.lower()
    content = await file.read()
    extracted_text = ""

    if filename.endswith(".pdf"):
        reader = PdfReader(BytesIO(content))
        for page in reader.pages:
            extracted_text += page.extract_text()
    elif filename.endswith(".txt"):
        extracted_text = content.decode("utf-8")
    elif filename.endswith(".doc") or filename.endswith(".docx"):
        doc = DocxDocument(BytesIO(content))
        for para in doc.paragraphs:
            extracted_text += para.text + "\n"
    else:
        raise HTTPException(status_code=400, detail="Unsupported file type")

    if not extracted_text.strip():
        raise HTTPException(status_code=400, detail="Could not extract text")

    response = await agent.run(f"User Query from Uploaded File:\n{extracted_text}")
    formatted_output = await generate_formatted_response(str(response.data))

    file_collection = chroma_client.get_or_create_collection(name="uploaded_files")

    file_collection.add(
        documents=[extracted_text],
        metadatas=[{
            "user_id": str(user.id),
            "filename": filename,
            "upload_date": datetime.now().isoformat(),
            "structured_data": json.dumps(response.data.dict())
        }],
        ids=[str(uuid.uuid4())]
    )

    for epic in response.data.epics:
        db_epic = Epic(
            user_id=user.id,
            title=epic.title,
            description=epic.description,
            estimation=epic.estimation
        )
        db.add(db_epic)
        db.flush()

        for story in epic.stories:
            db_story = Story(
                user_id=user.id,
                epic_id=db_epic.id,
                title=story.title,
                description=story.description,
                estimation=story.estimation
            )
            db.add(db_story)
            db.flush()

            for task in story.tasks:
                db_task = Task(
                    user_id=user.id,
                    story_id=db_story.id,
                    title=task.title,
                    description=task.description,
                    estimation=task.estimation
                )
                db.add(db_task)

    db.commit()

    return {
        "formatted_output": formatted_output,
        "structured_data": response.data
    }

@app.post("/search-similar-files")
async def search_similar_files(
    request: Request, 
    question: str = Form(...),
    db: Session = Depends(get_db) 
):

    user_email = request.session.get("user", {}).get("email")
    if not user_email:
        raise HTTPException(status_code=401, detail="User not logged in")
    
    user = db.query(User).filter(User.email == user_email).first()
    if not user:
        raise HTTPException(status_code=404, detail="User not found")

    file_collection = chroma_client.get_collection(name="uploaded_files")

    results = file_collection.query(
        query_texts=[question],
        n_results=3,
        where={"user_id": str(user.id)} 
    )
    return {
        "results": results
    }

@app.get("/download-conversation")
async def download_conversation(chat_id: int, format: str, db: Session = Depends(get_db)):
    conversations = db.query(Conversation).filter(Conversation.chat_id == chat_id).all()
    if not conversations:
        raise HTTPException(status_code=404, detail="Conversation not found")

    if format == "pdf":

        buffer = BytesIO()
        pdf = SimpleDocTemplate(buffer, pagesize=letter)
        elements = []

        styles = getSampleStyleSheet()
        for conv in conversations:
            elements.append(Paragraph(f"<b>User:</b> {conv.user_message}", styles["Normal"]))
            
            formatted_text = conv.bot_response.replace("\n", "<br/>")  
            elements.append(Paragraph(f"<b>Bot:</b><br/>{formatted_text}", styles["Normal"]))

            if conv.structured_data:
                structured_data = json.loads(conv.structured_data)
                if structured_data.get("epics"):
                    table_data = [["Type", "Title", "Description", "Estimation (Hours)"]]
                    
                    for epic in structured_data["epics"]:
                        table_data.append([
                            "Epic",
                            epic["title"],
                            epic["description"],
                            str(epic["estimation"])
                        ])
                        
                        for story in epic["stories"]:
                            table_data.append([
                                "Story",
                                story["title"],
                                story["description"],
                                str(story["estimation"])
                            ])
                            
                            for task in story["tasks"]:
                                table_data.append([
                                    "Task",
                                    task["title"],
                                    task["description"],
                                    str(task["estimation"])
                                ])

                    table = Table(table_data, colWidths=[80, 150, 250, 80])
                    table.setStyle(TableStyle([

                        ("BACKGROUND", (0, 0), (-1, 0), colors.HexColor("#4CAF50")),  
                        ("TEXTCOLOR", (0, 0), (-1, 0), colors.white),
                        ("FONTNAME", (0, 0), (-1, 0), "Helvetica-Bold"),
                        ("ALIGN", (0, 0), (-1, -1), "LEFT"),
                        ("FONTSIZE", (0, 0), (-1, 0), 12),
                        
                        ("BACKGROUND", (0, 1), (-1, 1), colors.HexColor("#187693")), 
                        ("TEXTCOLOR", (0, 1), (-1, 1), colors.white),

                        ("BACKGROUND", (0, 2), (-1, 2), colors.HexColor("#e81b1b66")), 
                        ("BACKGROUND", (0, 3), (-1, -1), colors.HexColor("#afdc1a3a")),  

                        ("GRID", (0, 0), (-1, -1), 1, colors.black),
                        ("BOX", (0, 0), (-1, -1), 1, colors.black),
                    ]))
                    elements.append(table)

        pdf.build(elements)
        buffer.seek(0)
        return StreamingResponse(
            buffer,
            media_type="application/pdf",
            headers={"Content-Disposition": "attachment; filename=conversation.pdf"}
        )
    if format == "word":
        document = DocxDocument()
        for conv in conversations:
            document.add_paragraph(f"User: {conv.user_message}")
            document.add_paragraph(f"Bot: {conv.bot_response}")

            if conv.structured_data:
                structured_data = json.loads(conv.structured_data)
                if structured_data.get("epics"):
                    
                    table = document.add_table(rows=1, cols=4)
                    table.style = "Table Grid"
                    hdr_cells = table.rows[0].cells
                    hdr_cells[0].text = "Type"
                    hdr_cells[1].text = "Title"
                    hdr_cells[2].text = "Description"
                    hdr_cells[3].text = "Estimation (Hours)"

                    for epic in structured_data["epics"]:
                        
                        row_cells = table.add_row().cells
                        row_cells[0].text = "Epic"
                        row_cells[1].text = epic["title"]
                        row_cells[2].text = epic["description"]
                        row_cells[3].text = str(epic["estimation"])

                        for story in epic["stories"]:
                            
                            row_cells = table.add_row().cells
                            row_cells[0].text = "Story"
                            row_cells[1].text = story["title"]
                            row_cells[2].text = story["description"]
                            row_cells[3].text = str(story["estimation"])

                            for task in story["tasks"]:
                                
                                row_cells = table.add_row().cells
                                row_cells[0].text = "Task"
                                row_cells[1].text = task["title"]
                                row_cells[2].text = task["description"]
                                row_cells[3].text = str(task["estimation"])

        
        buffer = BytesIO()
        document.save(buffer)
        buffer.seek(0)
        return StreamingResponse(
            buffer,
            media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            headers={"Content-Disposition": "attachment; filename=conversation.docx"}
        )
    else:
        raise HTTPException(status_code=400, detail="Invalid format")

@app.post("/store-conversation")
async def store_conversation(request: Request, db: Session = Depends(get_db)):
    form_data = await request.form()
    chat_id = form_data.get("chat_id")
    user_message = form_data.get("user_message")
    bot_response = form_data.get("bot_response")
    structured_data = form_data.get("structured_data")

    if not chat_id or not user_message or not bot_response:
        raise HTTPException(status_code=422, detail="Missing chat_id, user_message, or bot_response")

    existing_conversation = db.query(Conversation).filter(
        Conversation.chat_id == chat_id,
        Conversation.user_message == user_message,
        Conversation.bot_response == bot_response
    ).first()

    if existing_conversation:
        return JSONResponse({"status": "exists", "message": "Conversation already exists"})

    new_conversation = Conversation(
        chat_id=int(chat_id),
        user_message=user_message,
        bot_response=bot_response,
        structured_data=structured_data
    )
    db.add(new_conversation)
    db.commit()
    db.refresh(new_conversation)

    return JSONResponse({"status": "success", "message": "Conversation stored successfully"})

@app.get("/get-chat-name")
async def get_chat_name(request: Request, chat_id: int, db: Session = Depends(get_db)):
    chat = db.query(Chat).filter(Chat.id == chat_id).first()
    if not chat:
        raise HTTPException(status_code=404, detail="Chat not found")
    return JSONResponse({"status": "success", "chat_name": chat.chat_name})

@app.post("/submit-question")
async def submit_question(
    request: Request,
    question: str = Form(...),
    chat_id: int = Form(...),
    history: str = Form(...),
    db: Session = Depends(get_db)
):
    user_email = request.session.get("user", {}).get("email")
    user_id = None
    
    if user_email:
        user = db.query(User).filter(User.email == user_email).first()
        if user:
            user_id = user.id
    history_messages = json.loads(history) if history else []

    history_text = "\n".join([f"User: {msg['user']}\nBot: {msg['bot']}" for msg in history_messages])

    full_prompt = f"{history_text}\nUser: {question}\nBot:"

    results = collection.query(
        query_texts=[question],
        n_results=1
    )

    if results['documents'] and len(results['documents'][0]) > 0:
        existing_answer = results['documents'][0][0]
        refined_answer = await agent.run(f"Enhance the existing answer: {existing_answer} with the new question: {question}")
        answer = refined_answer.data
    else:
        result = await agent.run(full_prompt)
        answer = result.data

        collection.add(
            documents=[str(answer)],
            metadatas=[{
                "type": "epic_story_task",
                "user_id": str(user.id),
                "chat_id": str(chat_id),
                "timestamp": datetime.now().isoformat()
            }],
            ids=[str(uuid.uuid4())]
        )

    formatted_output = await generate_formatted_response(str(answer))

    new_conversation = Conversation(
        chat_id=int(chat_id),
        user_message=question,
        bot_response=formatted_output,
        structured_data=json.dumps(answer.dict())
    )
    db.add(new_conversation)
    db.commit()
    for epic_data in answer.epics:
        epic = Epic(
            title=epic_data.title,
            description=epic_data.description,
            estimation=epic_data.estimation,
            user_id=user_id  
        )
        db.add(epic)
        db.commit()
        
        for story_data in epic_data.stories:
            story = Story(
                title=story_data.title,
                description=story_data.description,
                estimation=story_data.estimation,
                epic_id=epic.id,
                user_id=user_id  
            )
            db.add(story)
            db.commit()
            
            for task_data in story_data.tasks:
                task = Task(
                    title=task_data.title,
                    description=task_data.description,
                    estimation=task_data.estimation,
                    story_id=story.id,
                    user_id=user_id  
                )
                db.add(task)
        db.commit()

    return JSONResponse({
        "formatted_output": formatted_output,
        "answer": answer.dict()
    })

@app.get("/get-conversations")
async def get_conversations(request: Request, chat_id: int, db: Session = Depends(get_db)):
    conversations = db.query(Conversation).filter(Conversation.chat_id == chat_id).distinct().all()

    conversation_list = []
    seen_messages = set()  

    for conv in conversations:
        key = (conv.user_message, conv.bot_response) 
        if key not in seen_messages:
            conversation_list.append({
                "user_message": conv.user_message,
                "bot_response": conv.bot_response,
                "structured_data": conv.structured_data
            })
            seen_messages.add(key)

    return JSONResponse({"status": "success", "conversations": conversation_list})


@app.get("/conversation.html", response_class=HTMLResponse)
async def conversation_page(request: Request):
    return templates.TemplateResponse("conversation.html", {"request": request})

@app.post("/create-chat")
async def create_chat(request: Request, chat_name: str = Form(...), db: Session = Depends(get_db)):
    user_email = request.session.get("user", {}).get("email")
    if not user_email:
        raise HTTPException(status_code=401, detail="User not logged in")

    user = db.query(User).filter(User.email == user_email).first()
    if not user:
        raise HTTPException(status_code=404, detail="User not found")

    new_chat = Chat(user_id=user.id, chat_name=chat_name)
    db.add(new_chat)
    db.commit()
    db.refresh(new_chat)

    return JSONResponse({"status": "success", "chat_id": new_chat.id})
@app.get("/get-chats")
async def get_chats(request: Request, db: Session = Depends(get_db)):
    user_email = request.session.get("user", {}).get("email")
    if not user_email:
        raise HTTPException(status_code=401, detail="User not logged in")

    user = db.query(User).filter(User.email == user_email).first()
    if not user:
        raise HTTPException(status_code=404, detail="User not found")

    chats = db.query(Chat).filter(Chat.user_id == user.id).all()
    chat_list = [{"id": chat.id, "chat_name": chat.chat_name} for chat in chats]

    return JSONResponse({"status": "success", "chats": chat_list})

@app.post("/delete-chat")
async def delete_chat(request: Request, chat_id: int = Form(...), db: Session = Depends(get_db)):
    user_email = request.session.get("user", {}).get("email")
    if not user_email:
        raise HTTPException(status_code=401, detail="User not logged in")

    user = db.query(User).filter(User.email == user_email).first()
    if not user:
        raise HTTPException(status_code=404, detail="User not found")

    chat = db.query(Chat).filter(Chat.id == chat_id, Chat.user_id == user.id).first()
    if not chat:
        raise HTTPException(status_code=404, detail="Chat not found")

    db.delete(chat)
    db.commit()

    return JSONResponse({"status": "success", "message": "Chat deleted successfully"})

@app.post("/update-chat")
async def update_chat(request: Request, db: Session = Depends(get_db)):
    form_data = await request.form()
    chat_id = form_data.get("chat_id")
    chat_name = form_data.get("chat_name")

    if not chat_id or not chat_name:
        raise HTTPException(status_code=422, detail="Missing chat_id or chat_name")

    user_email = request.session.get("user", {}).get("email")
    if not user_email:
        raise HTTPException(status_code=401, detail="User not logged in")

    user = db.query(User).filter(User.email == user_email).first()
    if not user:
        raise HTTPException(status_code=404, detail="User not found")

    chat = db.query(Chat).filter(Chat.id == int(chat_id), Chat.user_id == user.id).first()
    if not chat:
        raise HTTPException(status_code=404, detail="Chat not found")

    chat.chat_name = chat_name
    db.commit()

    return JSONResponse({"status": "success", "message": "Chat updated successfully", "chat_name": chat.chat_name})

@app.post("/verify-dob")
async def verify_dob(request: Request, db: Session = Depends(get_db)):
    data = await request.json()
    dob_verify = data.get("dob_verify")

    user_email = request.session.get("user", {}).get("email")
    if not user_email:
        return JSONResponse({"success": False, "message": "User not logged in"})

    user = db.query(User).filter(User.email == user_email).first()
    if user and user.dob == dob_verify:
        return JSONResponse({"success": True})
    else:
        return JSONResponse({"success": False, "message": "Invalid DOB"})

@app.get("/update-profile", response_class=HTMLResponse)
async def update_profile_page(request: Request, db: Session = Depends(get_db)):
    user_email = request.session.get("user", {}).get("email")
    if not user_email:
        return RedirectResponse(url="/onlylogin")
    user = db.query(User).filter(User.email == user_email).first()
    return templates.TemplateResponse("update_profile.html", {"request": request, "user": user})

@app.post("/update-profile")
async def update_profile(
    request: Request,
    first_name: str = Form(...),
    last_name: str = Form(...),
    country: str = Form(...),
    postal_code: str = Form(...),
    phone_number: str = Form(...),
    db: Session = Depends(get_db)
):
    user_email = request.session.get("user", {}).get("email")
    if not user_email:
        return RedirectResponse(url="/onlylogin")

    user = db.query(User).filter(User.email == user_email).first()
    if user:
        user.first_name = first_name
        user.last_name = last_name
        user.country = country
        user.postal_code = postal_code
        user.phone_number = phone_number
        db.commit()

    return RedirectResponse(url="/base", status_code=303)

@app.post("/update-password")
async def update_password(
    request: Request,
    new_password: str = Form(...),
    confirm_password: str = Form(...),
    db: Session = Depends(get_db)
):
    user_email = request.session.get("user", {}).get("email")       
    if not user_email:
        return RedirectResponse(url="/onlylogin")

    user = db.query(User).filter(User.email == user_email).first() 
    if user and new_password == confirm_password:
        user.password = pwd_context.hash(new_password)
        db.commit()
        return RedirectResponse(url="/base", status_code=303)
    else:
        raise HTTPException(status_code=400, detail="New password and confirm password do not match")

@app.post(")upload-profile-image")
async def upload_profile_image(request: Request, profile_image: UploadFile = File(...), db: Session = Depends(get_db)):
    user = request.session.get("user")
    if not user:
        raise HTTPException(status_code=401, detail="User not logged in")

    try:
        unique_filename = f"{user['email']}_{uuid.uuid4().hex}_{profile_image.filename}"
        file_path = os.path.join(UPLOAD_DIR, unique_filename)
        with open(file_path, "wb") as buffer:
            buffer.write(await profile_image.read())
        db_user = db.query(User).filter(User.email == user["email"]).first()
        if db_user:
            db_user.profile_image = file_path
            db.commit()
            db.refresh(db_user)
            request.session["user"]["profile_image"] = file_path

            return {"status": "success", "imagePath": f"/uploads/{unique_filename}"}
        else:
            raise HTTPException(status_code=404, detail="User not found")
    except Exception as e:
        print(f"Error uploading profile image: {e}")
        raise HTTPException(status_code=500, detail="An error occurred while updating the profile image")

@app.get("/", response_class=HTMLResponse)
async def home(request: Request):
    if "user" in request.session:
        request.session.clear()
    return templates.TemplateResponse("index.html", {"request": request})

@app.get("/base", response_class=HTMLResponse)
async def base_page(request: Request):
    user = request.session.get("user")
    return templates.TemplateResponse("base.html", {"request": request, "user": user})

@app.get("/onlylogin", response_class=HTMLResponse)
async def onlylogin_page(request: Request):
    return templates.TemplateResponse("onlylogin.html", {"request": request})

@app.get("/signup", response_class=HTMLResponse)
async def signup_page(request: Request):
    return templates.TemplateResponse("signup.html", {"request": request})

@app.api_route("/logout", methods=["GET", "POST"])
async def logout(request: Request):
    request.session.clear()
    response = RedirectResponse(url="/")
    response.delete_cookie("myapp_session")  
    return response

@app.post("/onlylogin")
async def login(request: Request, db: Session = Depends(get_db)):
    form_data = await request.form()
    email = form_data.get("email")
    password = form_data.get("password")

    user = db.query(User).filter(User.email == email).first()
    if user and pwd_context.verify(password, user.password):
        request.session["user"] = {
            "id": user.id,  # Make sure to include the user ID
            "email": user.email, 
            "name": user.first_name, 
            "profile_image": user.profile_image
        }
        return RedirectResponse(url="/base", status_code=303)
    
    raise HTTPException(status_code=401, detail="Invalid credentials")

@app.post("/signup")
async def signup(request: Request, db: Session = Depends(get_db)):
    form_data = await request.form()
    first_name = form_data.get("first_name")
    last_name = form_data.get("last_name")
    dob = form_data.get("dob")
    country = form_data.get("country")
    country_code = form_data.get("countryCode")
    phone_number = form_data.get("phone_number")
    gender = form_data.get("gender")
    email = form_data.get("email")
    password = form_data.get("password")
    
    if db.query(User).filter(User.email == email).first():
        raise HTTPException(status_code=400, detail="Email already registered")
    
    new_user = User(first_name=first_name, last_name=last_name, dob=dob, country=country, country_code=country_code, phone_number=phone_number, gender=gender, email=email, password=pwd_context.hash(password))
    db.add(new_user)
    db.commit()
    db.refresh(new_user)
    return RedirectResponse(url="/onlylogin", status_code=303)

message_history = []

@app.post("/submit", response_class=HTMLResponse)
async def submit(request: Request, question: str = Form(...)):
    try:
        db = SessionLocal()

        user = None
        user_email = request.session.get("user", {}).get("email")
        if user_email:
            user = db.query(User).filter(User.email == user_email).first()

        results = collection.query(
            query_texts=[question],
            n_results=1
        )

        if results['documents'] and len(results['documents'][0]) > 0:
            existing_answer = results['documents'][0][0]
            refined_answer = await agent.run(f"Enhance the existing answer: {existing_answer} with the new question: {question}")
            answer = refined_answer.data
        else:
            if not message_history:
                result = await agent.run(question)
                message_history.append({'question': question, 'response': result.data})
            else:
                prev_question = message_history[-1]['question']
                prev_response = message_history[-1]['response']
                combined_message = f"Modify this output based on the new question: {prev_response}\n\nNew Question: {question}"
                result = await agent.run(combined_message)
                message_history.append({'question': question, 'response': result.data})
            
            answer = message_history[-1]['response']

            collection.add(
                documents=[str(answer)],
                metadatas=[{"type": "epic_story_task"}],
                ids=[str(uuid.uuid4())]
            )

        for epic_data in answer.epics:
            epic = Epic(
                title=epic_data.title,
                description=epic_data.description,
                estimation=epic_data.estimation,
                user_id=user.id if user else None  
            )
            db.add(epic)
            db.commit()
            db.refresh(epic)
            
            for story_data in epic_data.stories:
                story = Story(
                    title=story_data.title,
                    description=story_data.description,
                    estimation=story_data.estimation,
                    epic_id=epic.id,
                    user_id=user.id if user else None  
                )
                db.add(story)
                db.commit()
                db.refresh(story)
                
                for task_data in story_data.tasks:
                    task = Task(
                        title=task_data.title,
                        description=task_data.description,
                        estimation=task_data.estimation,
                        story_id=story.id,
                        user_id=user.id if user else None
                    )
                    db.add(task)
            db.commit()

        formatted_output = await generate_formatted_response(str(answer))
        
        return templates.TemplateResponse("submit.html", {
            "request": request,
            "question": question,
            "formatted_output": formatted_output,
            "user": request.session.get("user"),
            "answer": answer 
        })

    except Exception as e:
        print(f"Error in submit endpoint: {str(e)}")
        raise HTTPException(status_code=500, detail=str(e))
    finally:
        db.close()